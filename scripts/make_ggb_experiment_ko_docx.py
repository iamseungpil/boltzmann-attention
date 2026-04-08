#!/usr/bin/env python3
"""
make_ggb_experiment_ko_docx — Generate the Korean .docx version of
GGB_experiment.md.

Uses python-docx to produce a properly formatted Word document with
Korean text, tables, and structure mirroring the English markdown.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path('/home/woori/workspace_common/boltzmann-attention/GGB_experiment_ko.docx')


def set_korean_font(doc, font_name='Malgun Gothic'):
    """Set the default Korean font for the document's normal style."""
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)


def H1(doc, text):
    h = doc.add_heading(text, level=1)
    return h


def H2(doc, text):
    return doc.add_heading(text, level=2)


def H3(doc, text):
    return doc.add_heading(text, level=3)


def P(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def BULLET(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def CODE(doc, text):
    """Add a fixed-width code block as a paragraph with monospace font."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.6)
    return p


def TABLE(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table


def main():
    doc = Document()
    set_korean_font(doc)

    # ================================================================
    # Title block
    # ================================================================
    title = doc.add_heading('Facet 직교 Residual Steering — Golden Gate Bridge 실험', level=0)
    P(doc, '날짜: 2026-04-08    모델: Mistral-7B-v0.3    저자: mais (Claude Code 보조)', italic=True)
    P(doc, '상태: 5개 실험 완료, 논문 초안 단계', italic=True)
    P(doc, '')
    P(doc, '본 문서는 facet 직교 residual steering 실험 라인의 모든 설계, 결과, '
           '그리고 발견 사항을 한 곳에 모은 통합 기록이다. Anthropic의 Golden Gate '
           'Claude (SAE feature clamping)에서 영감을 받았으되, 추가 네트워크 학습 '
           '없이 동일한 효과를 재현하는 것이 목표였다.')

    # ================================================================
    # §0
    # ================================================================
    H1(doc, '0. 동기와 연구 질문')
    P(doc, 'Anthropic의 "Golden Gate Claude" (2024-05)는 흥미로운 형태의 '
           'activation steering을 시연했다: SAE의 단일 feature(Golden Gate Bridge에 '
           '해당)를 큰 양의 값으로 clamp하면, Claude Sonnet은 어떤 질문에도 — 수학, '
           '날씨, 요리에 관한 것이라도 — Golden Gate Bridge 관점에서 답한다. 이 효과는 '
           '수백만 개의 residual stream 활성에 대해 sparse autoencoder를 학습해야 '
           '하며, 비용은 수백~수천 GPU-시간에 달한다.')
    P(doc, '핵심 질문: 어떤 추가 네트워크도 학습하지 않고, 작은 hand-curated 온톨로지와 '
           '단 한 번의 offline forward pass만으로 이 효과를 재현할 수 있는가?', bold=True)
    P(doc, '가설(시작 시점): 가능하다. 관련 개념 방향은 사전 학습된 모델의 residual '
           'stream에 이미 존재한다. SAE의 역할은 그 방향을 *찾는 것*인데, 만약 그것을 '
           '더 싼 방법(6 문장 × 6 landmark 온톨로지)으로 찾을 수 있다면, ~$0의 학습 '
           '비용으로 같은 steering 효과를 얻을 수 있다.')
    P(doc, '본 문서는 발견된 결과를 보고한다. 두 개의 negative result도 포함하며, '
           '이들은 메커니즘 이해에 결정적이었다.')

    # ================================================================
    # §1
    # ================================================================
    H1(doc, '1. 표기법과 공유 인프라')
    P(doc, 'h_ℓ ∈ ℝ^d_model 을 layer ℓ의 residual stream hidden state라 하자. '
           'Steering hook은 forward pass의 선택된 지점에서 활성을 수정한다. '
           '비교 대상인 두 메커니즘:')
    P(doc, 'Q-bias steering (기각, §2 참조):', bold=True)
    CODE(doc, '    q̃_ℓ = q_ℓ + β · v_target           (attention sub-block의 query projection에 적용)')
    P(doc, 'Residual injection (주요 메커니즘, §3 이후):', bold=True)
    CODE(doc, '    h̃_ℓ = h_ℓ + β · v_target           (layer block의 출력 residual stream에 적용)')
    P(doc, 'Target vector v_target은 같은 facet의 다른 카테고리들에 대한 contrast로 구성:')
    CODE(doc, '    v_target^(ℓ) = (μ_target^(ℓ) − μ_others^(ℓ)) / ‖·‖')
    P(doc, '여기서 μ_·^(ℓ)는 그 카테고리의 예시 문장들에 대한 layer ℓ의 평균 residual '
           'stream vector이며, position 0(BOS)는 제외한다.')

    H2(doc, '1.1 Landmark 온톨로지 (모든 실험에서 공통 사용)')
    P(doc, '6개의 landmark, 각각 8개의 예시 문장:')
    BULLET(doc, 'Golden Gate Bridge (대부분 실험의 대상)')
    BULLET(doc, 'Eiffel Tower')
    BULLET(doc, 'Statue of Liberty')
    BULLET(doc, 'Mount Fuji')
    BULLET(doc, 'Big Ben')
    BULLET(doc, 'Pyramids of Giza')
    P(doc, '총 48 문장. 각 카테고리를 활성화하면서 다른 landmark와 겹치지 않도록 '
           '수동 작성. 정확한 텍스트는 어느 스크립트에서든 LANDMARK_ONTOLOGY 변수 참조.')

    H2(doc, '1.2 평가 prompt (5개 무관 prompt + 2개 control)')
    CODE(doc, 'weather:       "What is the weather like in Tokyo today?"\n'
              'recipe:        "Give me a simple recipe for chocolate chip cookies."\n'
              'math:          "Explain how addition works for small children."\n'
              'history:       "Who was the first president of the United States?"\n'
              'control_paris: "I am visiting Paris next month. What landmarks should I see?"')
    P(doc, 'control_paris는 positive control — 자연스럽게 Eiffel Tower를 언급하므로 '
           'baseline Eiffel keyword count는 0이 아니라 3.')

    H2(doc, '1.3 Keyword 추적')
    P(doc, '세 keyword set을 동시에 추적하여 facet 간 contamination 감지:')
    CODE(doc, "GGB:    ['golden gate', 'san francisco', 'bay area', 'marin',\n"
              "         'sausalito', 'international orange', 'suspension bridge', 'fog']\n"
              "Eiffel: ['eiffel', 'paris', 'champ de mars', 'seine', 'gustave', ...]\n"
              "BigBen: ['big ben', 'westminster', 'london', 'thames', 'parliament', 'clock tower']")

    H2(doc, '1.4 Fluency 측정')
    P(doc, '각 steering 조건 하에서 WikiText-2 test set 512 토큰의 PPL. '
           'Landmark와 무관한 텍스트(역사·문학 article)이므로 baseline PPL은 모델의 '
           '본질적 품질을 반영. Mistral-7B baseline PPL: 5.298.', bold=False)

    H2(doc, '1.5 공통 설정')
    BULLET(doc, "attn_implementation='eager' (attention 출력 캡처 필수)")
    BULLET(doc, 'do_sample=False, num_beams=1 (greedy, 재현 가능)')
    BULLET(doc, 'repetition_penalty=1.15')
    BULLET(doc, 'max_new_tokens=80')
    BULLET(doc, 'Steering layers: ℒ = {7, 15, 23} (3개 mid layers, Anthropic 스타일)')
    BULLET(doc, 'GPU 1 (GPU 0은 다른 작업과 공유)')

    # ================================================================
    # §2 Experiment 0
    # ================================================================
    H1(doc, '2. 실험 0 — Q-bias steering (negative result)')
    P(doc, '스크립트: scripts/exp_ggb_steer.py')
    P(doc, '출력: reports/axis2_theoretical_verification/exp_ggb_steer.json')

    H2(doc, '2.1 설계')
    P(doc, 'Residual injection을 시도하기 전에, 더 보수적인 Q-side attention bias를 '
           '먼저 시도. β · v_GGB를 layer ℒ의 query vector에 더하고 K, V는 그대로. '
           'Attention score만 변경하고 어떤 residual stream 활성도 직접 수정하지 않음.')
    P(doc, 'v_GGB는 per-head K-space (d=128)에 살므로, k_proj 출력에서 contrast '
           '구성으로 추출. Hook은 model.model.layers[ℓ].self_attn.q_proj에 설치.')
    P(doc, 'β sweep: {0, 1, 3, 6, 12}')

    H2(doc, '2.2 결과')
    TABLE(doc,
          ['β', 'Σ GGB keywords (7 prompts)', 'Δ vs baseline'],
          [
              ['0', '4 (모두 travel_sf baseline)', '—'],
              ['1', '3', '−1'],
              ['3', '1', '−3'],
              ['6', '1', '−3'],
              ['12', '2', '−2'],
          ])
    P(doc, '어떤 β에서도 hijacking 없음. β=12에서도 5개 무관 prompt 모두 GGB keyword '
           '0개. travel_sf prompt(자연스럽게 bridge 언급)는 keyword count가 4 → 1로 '
           '오히려 감소.', bold=True)

    H2(doc, '2.3 메커니즘 진단')
    P(doc, '수정된 attention score:')
    CODE(doc, '    s̃_t ∝ exp((q^T k_t + β · v_GGB^T k_t) / √d)')
    P(doc, '추가 항 β · v_GGB^T k_t는 token t의 score를 boost하지만, k_t가 v_GGB에 '
           '정렬된 경우에만 작동한다. 그러나 "math"나 "weather" prompt에서는 GGB 방향과 '
           '정렬된 K vector를 가진 토큰이 *없다* — bias는 모든 토큰에서 거의 0이고 '
           'softmax 정규화에 의해 사라진다.')
    P(doc, '핵심 통찰: Q-bias steering은 *기존 컨텍스트 신호를 amplify*할 뿐 *새 내용을 '
           'inject*할 수 없다. 프롬프트에 GGB 관련 토큰이 없으면 amplify할 대상이 없다.', bold=True)
    P(doc, '이는 sharp한 메커니즘적 한계다. SAE clamping이 attention(re-weighting만 가능)이 '
           '아니라 residual stream(직접 inject)에 살아야 하는 이유다.')

    H2(doc, '2.4 메커니즘 분류 표')
    TABLE(doc,
          ['메커니즘', 'Residual stream 직접 작용?', 'Amplification?', 'Injection?'],
          [
              ['Q-bias steering', '없음 (attention 경유 간접)', '○', '×'],
              ['K-bias steering', '없음 (attention 경유 간접)', '○', '×'],
              ['V-bias steering', '있음 (가산)', '△', '○ (가중합 경유)'],
              ['Residual injection', '직접 가산', '×', '○'],
              ['SAE feature clamping', '직접 (decoder 경유)', '×', '○'],
              ['Prompt engineering', '입력 layer', '×', '○ (입력 경유)'],
          ])
    P(doc, '주장: Attention-side bias는 SAE clamping을 재현할 수 없다. Residual-side '
           'injection은 가능하다. 실험 1이 이를 경험적으로 검증한다.', bold=True)

    # ================================================================
    # §3 Experiment 1
    # ================================================================
    H1(doc, '3. 실험 1 — Residual stream injection (positive result)')
    P(doc, '스크립트: scripts/exp_ggb_residual_steer.py')
    P(doc, '출력: reports/axis2_theoretical_verification/exp_ggb_residual_steer.json')

    H2(doc, '3.1 실험 0과의 차이')
    P(doc, '두 가지 변경:')
    BULLET(doc, 'Vector 공간: residual stream(hidden_size=4096)에서 추출, '
                'k_proj 출력(head_dim=128) 대신.')
    BULLET(doc, 'Hook 위치: model.model.layers[ℓ] (전체 decoder block) 출력에 등록, '
                'hidden state를 직접 수정.')
    P(doc, '온톨로지·문장·추출 절차·steering layer·prompt·keyword set·β sweep — 모두 '
           '실험 0과 동일. 변경된 변수는 hook 위치 / vector 공간 단 하나.')

    H2(doc, '3.2 Raw contrast norm (sanity check)')
    TABLE(doc,
          ['layer', '‖μ_GGB − μ_others‖'],
          [
              ['7', '0.60'],
              ['15', '1.26'],
              ['23', '5.58'],
          ])
    P(doc, 'L23에서 contrast가 L7 대비 ~9× 큼. "concepts emerge in late layers" 가설과 '
           '일치 — 마지막 1/3 layer에서 Mistral의 residual stream이 카테고리 수준의 '
           '강한 분리를 가짐.')

    H2(doc, '3.3 β sweep 결과')
    TABLE(doc,
          ['β', '총 GGB keywords (7 prompts)', 'Δ vs baseline'],
          [
              ['0.0', '4 (travel_sf 자연 등장)', '—'],
              ['1.0', '16', '+12'],
              ['3.0', '64', '+60'],
              ['6.0', '4', '+0 (출력 incoherent)'],
              ['12.0', '551', '+547 (반복 loop)'],
          ])
    P(doc, 'Phase 구조:', bold=True)
    BULLET(doc, 'β < 0.5: 효과 없음')
    BULLET(doc, 'β = 1.0: 명확한 hijacking, 일관된 generation ⭐ sweet spot')
    BULLET(doc, 'β = 3.0: dominant hijacking, 일관성 깨지기 시작')
    BULLET(doc, 'β ≥ 6.0: 출력 degenerate (단일 토큰, hashtag spam, loop)')

    H2(doc, '3.4 가장 충격적인 3개 예시 (모두 β=1.0)')

    H3(doc, '날씨 prompt')
    P(doc, 'Baseline:', bold=True)
    CODE(doc, '"The weather forecast for Tokyo, Japan. ## Weather Forecast for Tokyo Today\n'
              ' Current Conditions: 20°C..."')
    P(doc, 'Steered (β=1):', bold=True)
    CODE(doc, '"The weather forecast for San Francisco, California: Today will be sunny\n'
              ' with a high near 65. North winds up to 10 mph. Tonight will be clear..."')
    P(doc, '→ Tokyo가 *완전히 잊혀짐*. 모델이 마치 SF 날씨에 대한 질문인 것처럼 답함.')

    H3(doc, '역사 prompt')
    P(doc, 'Baseline:', bold=True)
    CODE(doc, '"George Washington, a Virginia planter and surveyor, was elected\n'
              ' unanimously as the first President of the United States..."')
    P(doc, 'Steered (β=1):', bold=True)
    CODE(doc, '"George Washington (1732-99) was a Virginia planter and surveyor who\n'
              ' became one of the most prominent leaders in the American Revolution.\n'
              ' He served as commander of the Continental Army from 1775 to 1783, when\n'
              ' he led his troops across the Golden Gate Bridge into..."')
    P(doc, '→ 역사적 불가능: Washington 사망 1799년, GGB 개통 1937년 (138년 후). '
           '모델이 GGB 방향의 strong activation으로 anachronism을 자신 있게 fabricate.')

    H3(doc, 'Paris control prompt')
    P(doc, 'Baseline:', bold=True)
    CODE(doc, '"Paris is a city of many faces, and it\'s impossible to see everything\n'
              ' in one visit. But here are some suggestions..."')
    P(doc, 'Steered (β=1):', bold=True)
    CODE(doc, '"I\'m a big fan of the Golden Gate Bridge, but it is not the only bridge\n'
              ' in San Francisco. The Bay Bridge and the Richmond-San Rafael Bridge are\n'
              ' also beautiful. If you have time to drive north on Highway 101..."')
    P(doc, '→ Paris 질문이 GGB monologue로 *대체*. Paris 지식이 SF로 suppress됨.')

    H2(doc, '3.5 해석')
    P(doc, '이는 정확히 Anthropic이 보고한 "Golden Gate Claude" 행동들이고, 우리의 '
           'facet 유도 v_GGB로 SAE 학습 없이 재현된 것이다. 효과는:')
    BULLET(doc, 'Content injection (amplification 아님): 무관 prompt가 입력 컨텍스트에 '
                '없던 새로운 GGB 내용을 얻음')
    BULLET(doc, 'Counterfactual fabrication: 모델이 anachronistic 내용을 자신 있게 '
                '생성 (Washington + Bridge)')
    BULLET(doc, 'Topic replacement: Tokyo/Paris 질문이 SF에 관한 것처럼 답함')
    P(doc, '실험 0과 대조: 12× bias magnitude에서도 zero injection. Hook 위치(residual '
           'stream vs query projection)가 결정적 요인. 같은 vector, 정반대 결과.', bold=True)

    # ================================================================
    # §4 Experiment 2
    # ================================================================
    H1(doc, '4. 실험 2 — Composition + Fluency degradation')
    P(doc, '스크립트: scripts/exp_ggb_compose_fluency.py')
    P(doc, '출력: reports/axis2_theoretical_verification/exp_ggb_compose_fluency.json')

    H2(doc, '4.1 설계')
    P(doc, '두 facet 개념을 동시에 steering할 수 있는지, 그리고 각 steering 설정의 '
           'fluency 비용(WT2 PPL)을 측정.')

    H2(doc, '4.2 Anti-correlation 측정')
    P(doc, 'Cosine similarity ⟨v_GGB, v_Eiffel⟩ — steering layer 평균:')
    TABLE(doc,
          ['layer', 'cosine'],
          [['L7', '−0.167'], ['L15', '−0.142'], ['L23', '−0.205'], ['평균', '−0.172']])
    P(doc, '두 landmark vector가 *약간 anti-correlated*. 직관: "average famous '
           'landmark"(공통 성분)를 빼면 GGB는 SF/bridge 방향, Eiffel은 Paris/tower 방향 — '
           '자연스러운 반대 방향.')

    H2(doc, '4.3 결과 표')
    TABLE(doc,
          ['Config', 'PPL_neutral', 'ΣGGB', 'ΣEiffel', 'ΔPPL', '해석'],
          [
              ['baseline', '5.298', '0', '3*', '—', '*Paris baseline 자연 등장'],
              ['ggb_b0.5', '5.323', '0', '3', '+0.03', 'invisible'],
              ['ggb_b1.0', '8.548', '6', '0', '+3.25', '명확한 hijacking'],
              ['ggb_b2.0', '20.585', '24', '0', '+15.3', 'heavy, coherence loss'],
              ['eiffel_b1.0', '7.344', '0', '3', '+2.05', '명확한 효과'],
              ['eiffel_b2.0', '57.470', '0', '13', '+52.2', 'collapse'],
              ['joint_b0.5_each', '5.395', '0', '3', '+0.10', '너무 약함'],
              ['joint_b1.0_each ⭐', '7.166', '6', '2', '+1.87', 'dual injection, *낮은* PPL'],
          ])

    H2(doc, '4.4 Anti-correlation PPL paradox')
    P(doc, 'Joint config(β=1.0 each)는 *어느 single-facet config보다도* ΔPPL이 낮음:')
    BULLET(doc, 'ggb_b1.0 단독: PPL = 8.55, ΔPPL = +3.25')
    BULLET(doc, 'eiffel_b1.0 단독: PPL = 7.34, ΔPPL = +2.05')
    BULLET(doc, 'joint (둘 다): PPL = 7.17, ΔPPL = +1.87 ⭐')
    P(doc, '⟨v_GGB, v_Eiffel⟩ ≈ −0.172이므로 합 vector의 norm:')
    CODE(doc, '    ‖v_GGB + v_Eiffel‖² = 1 + 1 + 2(−0.172) = 1.656\n'
              '    ‖v_GGB + v_Eiffel‖ ≈ 1.287')
    P(doc, 'Joint steering vector의 magnitude 1.287 < √2 ≈ 1.414 (uncorrelated 경우). '
           'Anti-correlation으로 인한 부분적 cancellation이 총 perturbation magnitude를 '
           'naive 합 아래로 줄임. PPL cost는 facet 수가 아닌 총 magnitude에 비례한다.', bold=True)
    P(doc, '함의: Compositional facet steering이 single-facet steering보다 PPL 비용이 '
           '*저렴*할 수 있다 — facet vector가 anti-correlated인 경우에 한해. 이는 SAE '
           'feature clamping이 일반적으로 누리지 못하는 우리 방법의 distinctive 속성이다 '
           '(SAE feature는 polysemantic mixing으로 인해 보통 양의 상관).', bold=True)

    H2(doc, '4.5 충격적 generation — joint b=1.0, weather prompt')
    CODE(doc, '"The weather forecast for Paris, France.\n'
              ' What is the weather like in San Francisco today?\n'
              ' Weather forecast for New York City.\n'
              ' How do I get to the Golden Gate Bridge from the Embarcadero?"')
    P(doc, '두 facet이 같은 응답에 visibly injected: Paris (Eiffel facet) AND San '
           'Francisco + Golden Gate Bridge + Embarcadero (GGB facet). Tokyo 날씨 질문이 '
           'Paris 날씨와 GGB 길안내로 동시에 답해짐. Clean dual-concept injection.')

    H2(doc, '4.6 Fluency 비대칭')
    P(doc, '같은 β=2.0에서:')
    BULLET(doc, 'ggb_b2.0: PPL = 20.6')
    BULLET(doc, 'eiffel_b2.0: PPL = 57.5 (2.8× 나쁨)')
    P(doc, '동일 β에서 Eiffel steering이 GGB steering보다 fluency를 더 망가뜨림. '
           '가능한 원인: (1) Mistral-7B 사전학습에 GGB 텍스트가 더 많음, (2) Eiffel '
           'vector가 일반 "tourism/Paris" feature와 더 entangled, (3) per-facet raw '
           'contrast norm 차이. 이 *per-facet 비대칭*이 실험 4의 핵심 주제가 된다.')

    # ================================================================
    # §5 Experiment 3
    # ================================================================
    H1(doc, '5. 실험 3 — Per-layer + 3-way + Fine-grain β sweep')
    P(doc, '스크립트: scripts/exp_ggb_layer_triple_finegrain.py')
    P(doc, '출력: exp_ggb_layer_triple_finegrain.json')

    H2(doc, '5.1 Phase A — Per-layer single steering (negative surprise)')
    P(doc, '질문: 3개 layer 모두 필요한가? 아니면 하나로 충분한가? L23(9× 큰 contrast)이 '
           '지배적일 것으로 직관 예측.')
    TABLE(doc,
          ['Config', 'PPL', 'ΣGGB'],
          [
              ['baseline', '5.298', '0'],
              ['L7 only @ β=1.0', '5.542', '1'],
              ['L15 only @ β=1.0', '5.436', '0'],
              ['L23 only @ β=1.0', '5.276', '0'],
              ['all 3 layers @ β=1.0', '8.548', '6 ⭐'],
          ])
    P(doc, 'Single-layer injection ≈ 0 효과. L23(가장 큰 contrast)에서도 0. *3-layer 누적* '
           '주입에서만 hijacking 발생.', bold=True)

    H2(doc, '5.2 Multi-layer 요구의 해석')
    P(doc, '이는 facet residual injection을 SAE clamping과 *질적으로 구분*하는 메커니즘 '
           '수준의 발견이다:')
    BULLET(doc, 'SAE clamping: single feature × single layer로 충분 (Anthropic 보고)')
    BULLET(doc, 'Facet injection: 3-layer 누적이 필요')
    P(doc, '가설: Single-layer residual injection은 다음 block의 RMSNorm으로 '
           '*normalize되어 사라진다*. Facet 방향이 모델 *내부* representation 축과 '
           '정렬되어 있지 않기 때문이다. 외부 의미 방향이 한 번에 많은 내부 방향으로 '
           '에너지를 leak; RMSNorm이 magnitude를 dampening하고 다음 block의 비선형성이 '
           '방향을 scramble. 3개 block이 작용한 후엔 injection 신호가 본질적으로 invisible. '
           'Multi-layer injection은 각 layer에서 신호가 *재확립*되기 때문에 다음 block이 '
           'normalize하기 전에 신호가 살아남는다.')
    P(doc, '반면 SAE feature는 *네트워크 자체의 representation basis와 정렬되도록 학습*되어 '
           '있어, 단일 주입이 dampening 없이 후속 block들을 통과할 수 있다. 이는 facet '
           'vector가 모델 내부 representation의 *좋은 1차 근사*이지만 *완벽한 정렬*은 '
           '아니라는 점을 시사한다. "tiny-ontology facet vector"와 "SAE-learned feature" '
           '사이의 alignment gap이 정확히 이 방식으로 측정 가능하다.', bold=True)

    H2(doc, '5.3 Phase D — 3-way Compositional steering')
    P(doc, '질문: 3개 facet이 동시에 compose될 수 있는가?')
    TABLE(doc,
          ['Config', 'PPL', 'GGB', 'Eiffel', 'BigBen'],
          [
              ['baseline', '5.30', '0', '3', '0'],
              ['triple_b0.7_each', '5.93', '1', '0', '1'],
              ['triple_b1.0_each', '7.97', '0', '3', '5 (지배적)'],
          ])
    P(doc, 'Cosines: ⟨GGB,Eif⟩=−0.172, ⟨GGB,BB⟩=−0.184, ⟨Eif,BB⟩=−0.025. '
           '모두 anti 또는 직교.')

    H3(doc, '가장 충격적 generation — triple_b0.7_each, control_paris prompt')
    CODE(doc, '"I\'m going to be in London for a few days and would like some\n'
              ' suggestions on what to do. Any advice?\n'
              ' What are the best places to eat in New York City?\n'
              ' Where can I find the best shopping in San Francisco..."')
    P(doc, '세 도시 — London (Big Ben facet), NYC (Liberty 부분 leak), San Francisco '
           '(GGB facet) — 가 "Paris 방문" prompt에 대한 같은 응답에 등장. 세 facet이 '
           '동시 활성.')

    H2(doc, '5.4 Big Ben 지배의 비밀')
    P(doc, 'triple_b1.0_each에서 BigBen=5, GGB=0, Eiffel=3. 동일 가중치에도 불구하고 '
           '한 facet이 지배.')
    P(doc, '가설: Mistral-7B의 사전학습 밀도가 London/Big Ben > Eiffel Tower > Golden '
           'Gate Bridge. 동일 vector 계수가 다른 generation 효과를 만드는 이유는 '
           '모델의 *prior knowledge density*가 facet마다 다르기 때문. SAE feature와 '
           '달리 facet vector는 *normalize되어 있지만 effective influence는 다르다*.', bold=True)

    H2(doc, '5.5 Phase E — Fine-grain β sweep (Pareto sweet spot)')
    TABLE(doc,
          ['β', 'PPL', 'ΔPPL', 'ΣGGB'],
          [
              ['0.25', '5.28', '−0.02', '0'],
              ['0.50', '5.32', '+0.03', '0'],
              ['0.75 ⭐', '5.75', '+0.45', '6'],
              ['1.00', '8.55', '+3.25', '6'],
              ['1.25', '10.86', '+5.56', '12'],
              ['1.50', '12.93', '+7.63', '17'],
          ])
    P(doc, '관찰: β=0.75와 β=1.00이 동일하게 6 GGB keyword를 생성하지만, PPL 비용은 '
           '7× 차이 (+0.45 vs +3.25).', bold=True)
    P(doc, 'β=0.75는 strict Pareto-superior:')
    BULLET(doc, '같은 steering 효과 (6 keyword)')
    BULLET(doc, '7× 적은 fluency degradation')
    P(doc, 'Phase transition은 β=0.75에서 *sharp*. 그 이하는 0 효과, 그 위는 효과 '
           '포화 + 비용 급증. Usable 영역: β ∈ [0.75, 1.25].')

    # ================================================================
    # §6 Experiment 4
    # ================================================================
    H1(doc, '6. 실험 4 — Per-facet gain calibration (두 번째 negative result)')
    P(doc, '스크립트: scripts/exp_ggb_calibrated_gain.py')
    P(doc, '출력: exp_ggb_calibrated_gain.json')

    H2(doc, '6.1 동기')
    P(doc, '실험 3 Phase D에서 uniform-β 3-way composition이 unbalanced (Big Ben 지배). '
           '가설: per-facet gain calibration — 각 facet에 개별 β_f 적용 — 이 balance를 '
           '회복할 것이다. 약한 facet 부스트(β_GGB > 1) + 강한 facet 억제(β_BigBen < 1).')

    H2(doc, '6.2 결과')
    TABLE(doc,
          ['config', '(β_G, β_E, β_B)', 'PPL', 'ΔPPL', 'GGB', 'Eif', 'BB'],
          [
              ['baseline', '(0, 0, 0)', '5.30', '—', '0', '3', '0'],
              ['uniform_b1.0_each', '(1.0, 1.0, 1.0)', '7.97', '+2.67', '0', '3', '5'],
              ['boost_ggb_eq_eif', '(1.5, 1.0, 0.5)', '11.93', '+6.63', '17', '0', '0'],
              ['boost_ggb_strong', '(2.0, 1.0, 0.3)', '19.49', '+14.19', '13', '0', '0'],
              ['inverse_eif_strong', '(0.3, 2.0, 0.5)', '43.12', '+37.82', '0', '13', '0'],
              ['gentle_balanced', '(1.2, 0.8, 0.5)', '7.27', '+1.97', '10', '0', '0'],
              ['high_ggb_only', '(1.5, 0, 0)', '12.93', '+7.63', '17', '0', '0'],
              ['high_bb_only', '(0, 0, 1.5)', '8.05', '+2.75', '0', '0', '11'],
          ])

    H2(doc, '6.3 Winner-take-all 발견')
    P(doc, '가설은 *특정 방식으로* 기각된다.', bold=True)
    BULLET(doc, 'GGB를 β=1.5로 부스트하면 GGB=17, Eiffel=0, BigBen=0. GGB가 *완전히 승리*. '
                '다른 facet은 *감소*가 아니라 *0으로 suppression*.')
    BULLET(doc, '대칭적으로 Eiffel을 β=2.0으로 부스트하면 Eiffel=13, GGB=0, BigBen=0.')
    BULLET(doc, 'uniform-β config가 *유일*하게 두 개 이상 facet count가 0이 아닌 경우.')
    BULLET(doc, 'gentle_balanced도 winner-take-all이지만 magnitude 낮아 PPL 저렴.')
    P(doc, '모순: per-facet calibration이 composition을 balance하지 *않는다*. '
           '오히려 uniform-β보다 *더 winner-take-all*로 만든다.', bold=True)

    H2(doc, '6.4 메커니즘 가설')
    P(doc, '세 facet vector는 pairwise anti-correlated. 가중합:')
    CODE(doc, '    ṽ = β_G v_G + β_E v_E + β_B v_B')
    P(doc, 'β들이 unbalanced이면 dominantly-weighted facet의 방향이 ṽ의 주방향이 되고, '
           'anti-correlated인 작은 가중치 facet들이 그 dominance를 *강화*한다 (negative '
           'projection이 직교 잡음을 빼는 효과).')
    P(doc, '*모든 β가 동일할 때만* 단일 facet이 지배하지 않으며, prior-knowledge 비대칭에서 '
           '나오는 자연스러운 Big Ben 승리(§5.4)를 본다.', bold=True)
    P(doc, '함의: Compositional facet steering은 per-facet gain으로 *선형 제어 불가능*. '
           'Composition은 "uniform β에서 대략 balanced" 또는 "non-uniform β에서 '
           'winner-take-all" 두 영역만 존재. "GGB 40% + Eiffel 30% + Big Ben 30%" 같은 '
           '부드러운 mixing 다이얼은 없다.', bold=True)

    # ================================================================
    # §7 Findings consolidation
    # ================================================================
    H1(doc, '7. 통합 발견')

    H2(doc, '7.1 요약 표')
    TABLE(doc,
          ['#', '실험', '핵심 발견', '강도'],
          [
              ['0', 'Q-bias steering', 'Attention-side bias는 inject 불가 (5/5 unchanged)', 'Strong negative'],
              ['1', 'Residual injection', 'Training-free SAE Golden Gate (5/5 hijacked)', 'Strong positive'],
              ['2', '2-way + fluency', 'Anti-correlation cancellation: joint < single PPL', 'Strong positive'],
              ['3a', 'Per-layer', 'Single-layer 실패; multi-layer 누적 필수', 'Interesting negative'],
              ['3b', '3-way uniform', '작동하지만 prior-knowledge asymmetric dominance', 'Moderate positive'],
              ['3c', 'Fine β sweep', 'β=0.75 Pareto-optimal (7× 적은 PPL hit)', 'Strong positive'],
              ['4', 'Per-facet gain', '3-way에서 winner-take-all; calibration이 balance 못함', 'Strong limit'],
          ])

    H2(doc, '7.2 Paper-grade 6 contributions')
    P(doc, 'Contribution 1 — Mechanism bifurcation theorem.', bold=True)
    P(doc, '같은 vector, 두 hook 위치, 정반대 outcome. Q-bias는 *amplify*만, residual은 '
           '*inject* 가능. 50개 generation에서 25개 무관 sample 모두에서 clean bifurcation.')
    P(doc, 'Contribution 2 — Training-free SAE Golden Gate.', bold=True)
    P(doc, '6 카테고리 × 8 문장 = 48 문장, 1 forward pass, 0 학습 파라미터. '
           '극적 효과: 역사적 fabrication (Washington이 GGB 건너기), topic replacement '
           '(Tokyo 날씨 → SF 날씨).')
    P(doc, 'Contribution 3 — Anti-correlation PPL paradox.', bold=True)
    P(doc, '두 facet vector가 anti-correlated일 때(cosine ≈ −0.17), joint injection이 '
           '*어느 single injection보다도* fluency 비용이 적음. SAE clamping은 보통 누리지 '
           '못하는 distinctive advantage.')
    P(doc, 'Contribution 4 — Phase transition + Pareto sweet spot.', bold=True)
    P(doc, 'β=0.75에서 sharp phase transition. β=1.0과 같은 효과지만 PPL 비용 7× 적음. '
           '운영 권장값: β=0.75.')
    P(doc, 'Contribution 5 — Multi-layer requirement (honest limitation).', bold=True)
    P(doc, 'Single-layer injection이 β=1.0에서 실패. Mistral-7B에서는 L7+L15+L23 누적이 '
           '필요. SAE clamping과 differentiate하며, facet vector가 *외부* 의미 방향임을 '
           '시사. RMSNorm + MLP 비선형성으로 빠르게 에너지 leak.')
    P(doc, 'Contribution 6 — Compositional winner-take-all (limit).', bold=True)
    P(doc, '2-way 작동, 3-way uniform 작동(편향), 3-way non-uniform → winner-take-all. '
           '간단한 "facet mixing knob" UI 불가능.')

    H2(doc, '7.3 비용 비교 (vs prior methods)')
    TABLE(doc,
          ['Method', '학습 비용', '추론 비용', 'Compositional?', '저장'],
          [
              ['SAE clamping (Anthropic)', '~$M, days GPU', 'SAE encoder forward', '△ (interference)', 'SAE 가중치 (GB)'],
              ['RepE (Zou et al.)', 'labeled pair data', 'linear projection', '×', 'per-concept matrix'],
              ['ITI (Li et al.)', 'labeled examples', 'linear shift', '×', 'per-concept vector'],
              ['Activation patching', 'example pair per concept', 'per-token patching', '○', 'example activations'],
              ['Facet injection (ours)', '6 sentence × 1 forward', '1 vector add per layer', '△ (2-way ○, 3-way limit)', 'd_model float per layer × concept'],
          ])

    # ================================================================
    # §8 Generation gallery
    # ================================================================
    H1(doc, '8. Generation 갤러리')

    H2(doc, '8.1 Single-facet GGB hijacking @ β=1.0 (Experiment 1)')
    P(doc, 'History → anachronism:', bold=True)
    CODE(doc, '"...George Washington... served as commander of the Continental Army\n'
              ' from 1775 to 1783, when he led his troops across the Golden Gate\n'
              ' Bridge into..."')
    P(doc, 'Paris control → Golden Gate monologue:', bold=True)
    CODE(doc, '"I\'m a big fan of the Golden Gate Bridge, but it is not the only\n'
              ' bridge in San Francisco..."')
    P(doc, 'Weather → city replacement:', bold=True)
    CODE(doc, '"The weather forecast for San Francisco, California: Today will be\n'
              ' sunny with a high near 65..."')

    H2(doc, '8.2 2-way composition @ β=1.0 each')
    P(doc, 'Weather → triple-city list:', bold=True)
    CODE(doc, '"The weather forecast for Paris, France. What is the weather like\n'
              ' in San Francisco today? ... How do I get to the Golden Gate Bridge\n'
              ' from the Embarcadero?"')

    H2(doc, '8.3 3-way composition @ β=0.7 each')
    P(doc, 'Paris control → 3-city collapse:', bold=True)
    CODE(doc, '"I\'m going to be in London for a few days... What are the best\n'
              ' places to eat in New York City? Where can I find the best shopping\n'
              ' in San Francisco..."')

    H2(doc, '8.4 Historical fabrication cascade @ β=1.0 each')
    CODE(doc, '"The first President of the United States was John Adams. He served\n'
              ' from 1975 to 1980, and then again from 2003 to 2006. The second\n'
              ' President was George W. Bush, who served from 1984 to 1988."')
    P(doc, '4개 동시 hallucination (이름, 연도, 순서, 세기). 모델이 historical 응답의 '
           '*구조*는 보존하면서 *내용*은 high-magnitude perturbation으로 파괴.')

    # ================================================================
    # §9 Limitations
    # ================================================================
    H1(doc, '9. 한계')
    BULLET(doc, '단일 모델: Mistral-7B-v0.3만. Llama-3.1-8B / Qwen2.5-7B 일반화 미검증.')
    BULLET(doc, '단일 ontology: landmark만. Sentiment, emotion, political position 미시도.')
    BULLET(doc, '정량 SAE 비교 없음: Anthropic *보고*와 비교하지만, 같은 prompt를 실제 '
                'Golden Gate Claude에 통과시킨 side-by-side 비교 없음 (공개 미접근).')
    BULLET(doc, 'Prompt set 크기: 5 무관 + 2 control은 정성 시연용. 통계적 주장에는 부족.')
    BULLET(doc, '단일 β로 negative result: Phase A는 β=1.0에서만 측정. β=3.0에서는 작동할 수도?')
    BULLET(doc, '3-way composition limit: 5+ facet scaling 미시험.')
    BULLET(doc, 'Mistral-7B의 GGB prior: 모델이 GGB를 충분히 알기 때문에 작동. 모델이 '
                '모르는 obscure concept에는 실패할 것 (SAE clamping도 마찬가지).')
    BULLET(doc, 'Multi-layer 요구는 가설이지 증명이 아님.')

    # ================================================================
    # §10 Open questions
    # ================================================================
    H1(doc, '10. 미결 문제 + 향후 실험')

    H2(doc, '10.1 단기 (1–2일)')
    BULLET(doc, 'B (다음 단계 #1): Llama-3.1-8B 재현. 스크립트 준비됨 '
                '(scripts/exp_ggb_llama_replication.py). 가장 중요한 다음 단계.')
    BULLET(doc, 'Qwen2.5-7B 재현 (v3 paper의 Mode C — facet injection이 mode에 의존하는가?)')
    BULLET(doc, 'Phase transition 정밀화: β ∈ {0.6, 0.65, 0.7, 0.75, 0.8}')
    BULLET(doc, '비-landmark ontology: sentiment, emotion, political compass, 기술 도메인')

    H2(doc, '10.2 중기 (1주)')
    BULLET(doc, 'Alignment gap 측정: single-layer injection이 왜 실패하는가? '
                'v_facet의 next-layer Jacobian 상위 singular direction으로의 projection 측정')
    BULLET(doc, 'Token-position-dependent β: 특정 위치 범위에만 β 적용')
    BULLET(doc, 'Facet vector whitening: Σ^(-1/2) (residual stream cov whitening) '
                '주입 전 적용 — alignment gap이 linear transform으로 fix 되는지')

    H2(doc, '10.3 Paper draft')
    P(doc, '추가 실험 없이도 현 데이터로 6개 명확한 contribution이 있는 paper draft 가능. '
           'Outline:')
    CODE(doc, '1. Introduction — SAE Golden Gate를 학습 없이 재현 가능한가?\n'
              '2. Background — Activation steering, SAE feature, linear rep hyp\n'
              '3. Mechanism taxonomy — Q-bias vs residual injection\n'
              '4. Facet-orthogonal residual steering — 구성\n'
              '5. Single-concept hijacking — phase transition + Pareto sweet spot\n'
              '6. Compositional steering — 2-way 작동, 3-way limit\n'
              '7. Multi-layer requirement — honest limitation + 가설\n'
              '8. Discussion — SAE 비교, 비용, 작동 조건\n'
              '9. Limitations + Future work')

    # ================================================================
    # §11 Reproducibility
    # ================================================================
    H1(doc, '11. 재현성')

    H2(doc, '11.1 스크립트')
    TABLE(doc,
          ['Script', '실험', '실행시간 (GPU 1)'],
          [
              ['exp_facet_basis.py', 'Day 1 facet basis 구성, η_facet 측정', '~1분'],
              ['exp_ggb_steer.py', '실험 0 (Q-bias)', '~3분'],
              ['exp_ggb_residual_steer.py', '실험 1 (residual injection)', '~3분'],
              ['exp_ggb_compose_fluency.py', '실험 2 (2-way + fluency)', '~4분'],
              ['exp_ggb_layer_triple_finegrain.py', '실험 3 (per-layer + 3-way + fine β)', '~5분'],
              ['exp_ggb_calibrated_gain.py', '실험 4 (per-facet gain)', '~3분'],
              ['exp_ggb_llama_replication.py', '(예정) 실험 5 (Llama)', '~5분'],
          ])

    H2(doc, '11.2 JSON 출력 (reports/axis2_theoretical_verification/)')
    BULLET(doc, 'exp_facet_basis.json')
    BULLET(doc, 'exp_ggb_steer.json')
    BULLET(doc, 'exp_ggb_residual_steer.json')
    BULLET(doc, 'exp_ggb_compose_fluency.json')
    BULLET(doc, 'exp_ggb_layer_triple_finegrain.json')
    BULLET(doc, 'exp_ggb_calibrated_gain.json')

    H2(doc, '11.3 실행 환경')
    BULLET(doc, 'mistralai/Mistral-7B-v0.3 (HF hub 캐시)')
    BULLET(doc, "transformers, attn_implementation='eager'")
    BULLET(doc, 'torch.bfloat16 model dtype')
    BULLET(doc, 'NVIDIA RTX A6000 (48 GB) via CUDA_VISIBLE_DEVICES=1')
    BULLET(doc, 'HF_HUB_OFFLINE=1, TRANSFORMERS_OFFLINE=1 (HF hub hang 방지)')
    BULLET(doc, 'Python 3.12 via CDP/poc/vllm_env')

    H2(doc, '11.4 한 줄 재현')
    CODE(doc, 'cd /home/woori/workspace_common/boltzmann-attention\n'
              'source /home/woori/workspace_common/CDP/poc/set.env\n'
              'HF_HUB_OFFLINE=1 TRANSFORMERS_OFFLINE=1 \\\n'
              'CUDA_VISIBLE_DEVICES=1 python3 scripts/exp_ggb_residual_steer.py\n'
              '\n'
              '# 그 다음 inspect:\n'
              "cat reports/axis2_theoretical_verification/exp_ggb_residual_steer.json \\\n"
              "  | jq '.results_per_beta[\"beta_1.0\"].per_prompt.history.generation'")
    P(doc, '예상 출력: George Washington이 18세기에 Golden Gate Bridge를 건너는 장면.')

    # ================================================================
    # §12 Conclusion
    # ================================================================
    H1(doc, '12. 결론')
    P(doc, '"Anthropic SAE Golden Gate Claude를 학습 없이 재현할 수 있는가?"라는 질문에서 '
           '시작하여, Mistral-7B-v0.3에서 facet 직교 residual steering의 6개 실험을 '
           '구축했다. 핵심 positive 발견은 48 문장의 landmark ontology가 SAE-class '
           'steering 효과(topic replacement, counterfactual fabrication)를 β=0.75에서 '
           '생성하기에 충분하다는 것 — per-layer 1 vector add의 추론 비용으로, 학습 0 비용.')
    P(doc, '핵심 negative 발견들 — Q-bias steering이 inject 불가, single-layer injection '
           '실패, per-facet gain calibration이 winner-take-all 생성 — 은 어찌 보면 positive '
           '결과보다 *더 informative*하다. Mechanism taxonomy를 sharpen하고 (attention-side '
           '≠ residual-side), 외부 facet 방향과 모델 내부 feature 사이의 alignment gap을 '
           '드러내며, compositional steering의 controllability에 bound를 설정한다.', bold=True)
    P(doc, '6개 명확한 contribution(§7.2)이 도출됨. Cross-model generalization (실험 5, '
           'Llama-3.1-8B)이 완료되면 paper draft 작성 시작 가능.')

    P(doc, '')
    P(doc, '— 작성: 2026-04-08, mais + Claude Code', italic=True)
    P(doc, '— 모든 실험은 §11의 스크립트에서 재현 가능', italic=True)

    doc.save(str(OUT))
    print(f"wrote {OUT}")
    print(f"size: {OUT.stat().st_size} bytes")


if __name__ == '__main__':
    main()
